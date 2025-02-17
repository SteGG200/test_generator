from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, ns
from docx.oxml.xmlchemy import BaseOxmlElement
from docx.text.run import Run

def create_element(name):
	return OxmlElement(name)

def create_attribute(element : BaseOxmlElement, name : str, value: str):
	element.set(ns.qn(name), value)


def add_page_number(run: Run):
	page_element = create_element('w:t')
	create_attribute(page_element, 'xml:space', 'preserve')
	page_element.text = 'Trang '

	run._r.append(page_element)

	fldChar1 = create_element('w:fldChar')
	create_attribute(fldChar1, 'w:fldCharType', 'begin')

	instrText1 = create_element('w:instrText')
	create_attribute(instrText1, 'xml:space', 'preserve')
	instrText1.text = "PAGE"

	fldChar2 = create_element('w:fldChar')
	create_attribute(fldChar2, 'w:fldCharType', 'end')

	run._r.append(fldChar1)
	run._r.append(instrText1)
	run._r.append(fldChar2)

	slash = create_element('w:t')
	create_attribute(slash, 'xml:space', 'preserve')
	slash.text = ' / '

	run._r.append(slash)

	fldChar3 = create_element('w:fldChar')
	create_attribute(fldChar3, 'w:fldCharType', 'begin')

	instrText2 = create_element('w:instrText')
	create_attribute(instrText2, 'xml:space', 'preserve')
	instrText2.text = "NUMPAGES"

	fldChar4 = create_element('w:fldChar')
	create_attribute(fldChar4, 'w:fldCharType', 'end')

	run._r.append(fldChar3)
	run._r.append(instrText2)
	run._r.append(fldChar4)

def create_exam_document(exam_content: str):
	if len(exam_content) == 0:
		raise ValueError("Nội dung đề thi không hợp lệ")

	doc = Document()

	# Add page numbers and footers
	for section in doc.sections:
		section.left_margin = Inches(0.9)
		section.right_margin = Inches(0.6)
		footer = section.footer
		footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
		add_page_number(footer_para.add_run())
		footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

	# Style
	style = doc.styles['Normal']
	font = style.font
	font.name = 'Times New Roman'
	font.size = Pt(14)
	
	table = doc.add_table(rows=1, cols=2)

	# Header
	left = table.cell(0, 0)
	left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
	header = left.add_paragraph()
	header.alignment = WD_ALIGN_PARAGRAPH.CENTER
	header.add_run('TRƯỜNG ĐẠI HỌC KHOA HỌC TỰ NHIÊN\n').font.size = Pt(12)
	header.add_run('TRƯỜNG THPT CHUYÊN KHTN\n').bold = True
	
	title = left.add_paragraph()
	title.alignment = WD_ALIGN_PARAGRAPH.CENTER
	title.add_run('ĐỀ LUYỆN TẬP\n\n').bold = True
	
	# Exam info
	right = table.cell(0, 1)
	right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
	info = right.add_paragraph()
	info.alignment = WD_ALIGN_PARAGRAPH.CENTER
	info.add_run('ĐỀ THI ÔN TẬP CHỦ ĐỀ\n').bold = True
	info.add_run('Môn: Tin học\n').bold = True
	info.add_run('Thời gian làm bài: ... phút, không kể thời gian phát đề\n\n').italic = True
	
	# Student info
	student_info = doc.add_paragraph()
	student_info.add_run('\nHọ tên thí sinh:.....................................................................').bold = True
	student_info.add_run('Số báo danh:........................\n').bold = True
	
	# Instructions
	instructions = doc.add_paragraph()
	instructions.add_run('PHẦN I. Câu trắc nghiệm nhiều phương án lựa chọn. ').bold = True
	
	# Add exam content
	lines = exam_content.split('\n')
	total_number_questions = 0
	question_section = None
	for line in lines:
		if line == '': continue
		if line.startswith("**Câu"):
			question_section = doc.add_paragraph()
			total_number_questions += 1
			number_question, content_question = line[2:].split('**')
			question_section.add_run(number_question).bold = True
			question_section.add_run(content_question)
		elif question_section == None: raise ValueError("Invalid question format.")
		elif line.strip().startswith(('**A', '**B', '**C', '**D')):
			order_answer, content_answer = line[2:].split('**')
			question_section.add_run(f'\n{order_answer}').bold = True
			question_section.add_run(f'{content_answer}')
		else:
			question_section.add_run(f'\n{line}')

	# Add total number of questions
	instructions.add_run(f'Thí sinh trả lời từ câu 1 đến câu {total_number_questions}. Mỗi câu hỏi thí sinh chỉ chọn một phương án.')
	
	# Footer note
	footer = doc.add_paragraph()
	footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
	footer.add_run('\n\n----------------------- HẾT -----------------------\n\n').bold = True
	footer.add_run('- Thí sinh không được sử dụng tài liệu;\n').italic = True
	footer.add_run('- Cán bộ coi thi không giải thích gì thêm.').italic = True
	
	return doc