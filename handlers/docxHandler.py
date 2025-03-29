import os

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns
from docx.oxml.xmlchemy import BaseOxmlElement
from docx.shared import Inches, Pt
from docx.text import paragraph
from docx.text.run import Run
from rich import print


def create_element(name):
    return OxmlElement(name)


def create_attribute(element: BaseOxmlElement, name: str, value: str):
    element.set(ns.qn(name), value)


def add_page_number(run: Run):
    page_element = create_element("w:t")
    create_attribute(page_element, "xml:space", "preserve")
    page_element.text = "Trang "

    run._r.append(page_element)

    fldChar1 = create_element("w:fldChar")
    create_attribute(fldChar1, "w:fldCharType", "begin")

    instrText1 = create_element("w:instrText")
    create_attribute(instrText1, "xml:space", "preserve")
    instrText1.text = "PAGE"

    fldChar2 = create_element("w:fldChar")
    create_attribute(fldChar2, "w:fldCharType", "end")

    run._r.append(fldChar1)
    run._r.append(instrText1)
    run._r.append(fldChar2)

    slash = create_element("w:t")
    create_attribute(slash, "xml:space", "preserve")
    slash.text = " / "

    run._r.append(slash)

    fldChar3 = create_element("w:fldChar")
    create_attribute(fldChar3, "w:fldCharType", "begin")

    instrText2 = create_element("w:instrText")
    create_attribute(instrText2, "xml:space", "preserve")
    instrText2.text = "NUMPAGES"

    fldChar4 = create_element("w:fldChar")
    create_attribute(fldChar4, "w:fldCharType", "end")

    run._r.append(fldChar3)
    run._r.append(instrText2)
    run._r.append(fldChar4)


def line_detect(line: str):
    """
    Check if line is in question or option content

    0 - Child line (for multi-line questions or options)
    1 - First line of question
    2 - First line of option
    3 - First line of correct option
    -1 - Invalid line (not a question or option)
    """
    if "\n" in line:
        raise ValueError(f'"{line}" is not a single line')
    if line == "" or line.lstrip() != line:
        return 0
    elif line[0].isnumeric():
        return 1
    elif len(line) >= 2 and line[0].isalpha() and line[1] == ")":
        return 2
    elif len(line) >= 3 and line[0] == "*" and line[1].isalpha() and line[2] == ")":
        return 3
    else:
        return -1


def docx_handler(path: str):
    with open(os.path.join(path, "content.txt")) as log:
        exam_content = log.read()

    if len(exam_content) == 0:
        raise ValueError("Nội dung đề thi không hợp lệ")

    doc = Document()

    # Add page numbers and footers
    for section in doc.sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.6)
        footer = section.footer
        footer_para = (
            footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        )
        add_page_number(footer_para.add_run())
        footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Style
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(14)

    table = doc.add_table(rows=1, cols=2)

    # Header
    left = table.cell(0, 0)
    left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    header = left.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.add_run("TRƯỜNG ĐẠI HỌC KHOA HỌC TỰ NHIÊN\n").font.size = Pt(12)
    header.add_run("TRƯỜNG THPT CHUYÊN KHTN\n").bold = True

    title = left.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("ĐỀ LUYỆN TẬP\n\n").bold = True

    # Exam info
    right = table.cell(0, 1)
    right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    info = right.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("ĐỀ THI ÔN TẬP CHỦ ĐỀ\n").bold = True
    info.add_run("Môn: Tin học\n").bold = True
    info.add_run(
        "Thời gian làm bài: ... phút, không kể thời gian phát đề\n\n"
    ).italic = True

    # Student info
    student_info = doc.add_paragraph()
    student_info.add_run(
        "\nHọ tên thí sinh:....................................................................."
    ).bold = True
    student_info.add_run("Số báo danh:........................\n").bold = True

    # Instructions
    instructions = doc.add_paragraph()
    instructions.add_run(
        "PHẦN I. Câu trắc nghiệm nhiều phương án lựa chọn. "
    ).bold = True

    # Add exam content
    lines = exam_content.split("\n")
    total_number_questions = 0
    current_line_number = 1
    question_paragraph: paragraph.Paragraph | None = None
    while current_line_number <= len(lines):
        current_line = lines[current_line_number - 1]
        type_current_line = line_detect(current_line)
        if type_current_line == -1:
            raise ValueError(
                f'In "{path}/{"content.txt"}" on line {current_line_number}:\nUnexpected line'
            )
        elif type_current_line == 0:
            if len(current_line.strip()) != 0:
                raise ValueError(
                    f'In "{path}/{"content.txt"}" on line {current_line_number}:\nInvalid content format'
                )
            current_line_number += 1
        else:
            content_lines: list[str] = [current_line]
            buffer_line_number = current_line_number + 1
            while (
                buffer_line_number <= len(lines)
                and line_detect(lines[buffer_line_number - 1]) == 0
            ):
                content_lines.append(lines[buffer_line_number - 1])
                buffer_line_number += 1

            while content_lines[-1].strip() == "":
                content_lines.pop()

            if type_current_line == 2:  # Normal option
                for index, content_line in enumerate(content_lines):
                    if line_detect(content_line) == 0:
                        question_paragraph.add_run(f"\n{content_line}")
                    else:
                        array_option = content_line.split(")", 1)
                        if len(array_option) != 2:
                            raise ValueError(
                                f'In "{path}/{"content.txt"}" on line {current_line_number + index}:\nInvalid option format'
                            )
                        order_option, content_option = array_option
                        question_paragraph.add_run(
                            f"\n{order_option.upper()}. "
                        ).bold = True
                        question_paragraph.add_run(content_option)
                if (
                    buffer_line_number <= len(lines)
                    and line_detect(lines[buffer_line_number - 1]) == 1
                ):
                    # Set None at the end of question paragraph
                    question_paragraph = None
            elif type_current_line == 3:  # Correct option
                for index, content_line in enumerate(content_lines):
                    if line_detect(content_line) == 0:
                        question_paragraph.add_run(f"\n{content_line}")
                    else:
                        array_option = content_line[1:].split(")", 1)
                        if len(array_option) != 2:
                            raise ValueError(
                                f'In "{path}/{"content.txt"}" on line {current_line_number + index}:\nInvalid correct option format'
                            )
                        order_option, content_option = array_option
                        order_option_runner = question_paragraph.add_run(
                            f"\n{order_option.upper()}. "
                        )
                        order_option_runner.underline = True
                        order_option_runner.bold = True
                        question_paragraph.add_run(content_option)
                if (
                    buffer_line_number <= len(lines)
                    and line_detect(lines[buffer_line_number - 1]) == 1
                ):
                    # Set None at the end of question paragraph
                    question_paragraph = None
            else:  # Question content
                for index, content_line in enumerate(content_lines):
                    if len(content_line) > 0 and content_line[0].isnumeric():
                        index_separator = content_line.find(".")
                        if index_separator == -1:
                            raise ValueError(
                                f'In "{path}/{"content.txt"}" on line {current_line_number + index}:\nInvalid question content format'
                            )
                        array_answer = content_line.split(".", 1)
                        if len(array_answer) != 2:
                            raise ValueError(
                                f'In "{path}/{"content.txt"}" on line {current_line_number + index}:\nInvalid question content format'
                            )
                        number_question, content_question = array_answer
                        if not number_question.isnumeric():
                            raise ValueError(
                                f'In "{path}/{"content.txt"}" on line {current_line_number + index}:\nInvalid question content format'
                            )

                        if question_paragraph is not None:
                            raise ValueError(
                                f'In "{path}/{"content.txt"} on line {current_line_number + index}:\nInvalid question content format'
                            )

                        question_paragraph = doc.add_paragraph()
                        total_number_questions += 1
                        question_paragraph.add_run(
                            f"Câu {number_question}: "
                        ).bold = True
                        question_paragraph.add_run(content_question)
                    else:
                        question_paragraph.add_run(f"\n{content_line}")
            current_line_number = buffer_line_number

    # Add total number of questions
    instructions.add_run(
        f"Thí sinh trả lời từ câu 1 đến câu {total_number_questions}. Mỗi câu hỏi thí sinh chỉ chọn một phương án."
    )

    # Footer note
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(
        "\n\n----------------------- HẾT -----------------------\n\n"
    ).bold = True
    footer.add_run("- Thí sinh không được sử dụng tài liệu;\n").italic = True
    footer.add_run("- Cán bộ coi thi không giải thích gì thêm.").italic = True

    doc.save(os.path.join(path, DOCX_FILE))

    print(
        "[blue]    └── [/blue]"
        "[green]Đã tạo file docx thành công: [/green]"
        f"[white]{os.path.join(path, DOCX_FILE)}[/white]"
    )
